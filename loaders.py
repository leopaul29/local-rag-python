"""Phase 1 - turn files on disk into raw text units.

Each loader yields RawDocument objects. Keeping the source and the page
number here is what makes citations possible at the very end of the chain.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

# CJK punctuation, kana, kanji and full-width forms
_CJK_CHAR = (
    r"[\u3000-\u303f\u3040-\u309f\u30a0-\u30ff"
    r"\u3400-\u4dbf\u4e00-\u9fff\uff00-\uffef]"
)


@dataclass
class RawDocument:
    """One logical piece of a file: a PDF page, or a whole text file."""

    source: str  # file name, shown to the user in citations
    page: int | None  # 1-based page number, None for non-paginated formats
    text: str


def _clean(text: str) -> str:
    """Normalize whitespace without destroying paragraph boundaries."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse 3+ newlines into a clean paragraph separator
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Repair words split across line breaks by PDF extraction ("exam-\nple")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    # Japanese PDFs hard-wrap mid-sentence and the break carries no meaning.
    # Join a single newline when CJK text sits on both sides. Blank lines are
    # untouched, so paragraph boundaries survive.
    # Numbers and latin words are common inside Japanese sentences
    # (合計は\n125,000円), so accept them on either side of the break.
    text = re.sub(rf"(?<={_CJK_CHAR})\n(?={_CJK_CHAR}|[0-9A-Za-z])", "", text)
    text = re.sub(rf"(?<=[0-9A-Za-z])\n(?={_CJK_CHAR})", "", text)
    return text.strip()


def load_pdf(path: Path) -> Iterator[RawDocument]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ImportError("PDF support requires 'pypdf' (pip install pypdf)") from exc

    reader = PdfReader(str(path))
    for index, page in enumerate(reader.pages, start=1):
        text = _clean(page.extract_text() or "")
        if text:
            yield RawDocument(source=path.name, page=index, text=text)


def load_docx(path: Path) -> Iterator[RawDocument]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ImportError("DOCX support requires 'python-docx'") from exc

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]

    # Tables often hold the most valuable facts, so flatten them too
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = _clean("\n\n".join(p for p in parts if p.strip()))
    if text:
        yield RawDocument(source=path.name, page=None, text=text)


def load_text(path: Path) -> Iterator[RawDocument]:
    text = _clean(path.read_text(encoding="utf-8", errors="replace"))
    if text:
        yield RawDocument(source=path.name, page=None, text=text)


LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_text,
    ".md": load_text,
    ".markdown": load_text,
}


def load_directory(directory: Path) -> Iterator[RawDocument]:
    """Walk a directory recursively and yield every readable document."""
    if not directory.exists():
        raise FileNotFoundError(f"Documents directory not found: {directory}")

    for path in sorted(directory.rglob("*")):
        if not path.is_file() or path.name.startswith("."):
            continue

        loader = LOADERS.get(path.suffix.lower())
        if loader is None:
            print(f"  skipped (unsupported format): {path.name}")
            continue

        try:
            yield from loader(path)
        except Exception as exc:  # keep going, one broken file is not fatal
            print(f"  failed to read {path.name}: {exc}")
